"""
Specific File Processors for Document Ingestion

This module provides concrete implementations of BaseDocumentProcessor for different
file formats including PDF, Word documents, text files, and web pages.

Key Processors:
- PDFProcessor: Extract text and metadata from PDF files
- WordProcessor: Handle .docx and .doc files  
- TextProcessor: Process plain text files (.txt, .md)
- WebProcessor: Extract content from web pages and HTML files
- DocumentProcessorFactory: Factory for creating appropriate processors
"""

import io
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging
import requests
from urllib.parse import urlparse, urljoin
import time
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

# Import our base classes
from .document_ingestion import (
    BaseDocumentProcessor, ProcessingError, 
    UnsupportedFormatError
)

# Configure logging
logger = logging.getLogger(__name__)


class PDFProcessor(BaseDocumentProcessor):
    """
    Process PDF documents using PyPDF2.
    
    Features:
    - Extract text content from all pages
    - Extract PDF metadata (title, author, creator, etc.)
    - Handle password-protected PDFs
    - Page-by-page processing for large files
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.password = self.config.get('password', None)
    
    def _get_supported_extensions(self) -> List[str]:
        return ['.pdf']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Handle password-protected PDFs
                if pdf_reader.is_encrypted:
                    if self.password:
                        pdf_reader.decrypt(self.password)
                    else:
                        raise ProcessingError(f"PDF is password-protected: {file_path}")
                
                # Extract text from all pages with enhanced metadata
                text_content = []
                page_metadata = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            # Store page metadata for later citation use
                            page_info = {
                                'page_number': page_num + 1,
                                'text_length': len(page_text.strip()),
                                'has_images': '/XObject' in page.get('/Resources', {}),
                            }
                            page_metadata.append(page_info)
                            
                            # Add page marker with metadata for chunking
                            page_header = f"\\n--- Page {page_num + 1} ---\\n"
                            text_content.append(f"{page_header}{page_text}")
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num + 1} in {file_path}: {e}")
                        continue
                
                if not text_content:
                    raise ProcessingError(f"No extractable text found in PDF: {file_path}")
                
                return "\\n".join(text_content)
                
        except Exception as e:
            if isinstance(e, ProcessingError):
                raise
            raise ProcessingError(f"Failed to extract text from PDF {file_path}: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'processor_type': 'pdf',
                    'page_count': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted
                }
                
                # Extract PDF metadata if available
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata.update({
                        'title': getattr(pdf_info, 'title', None),
                        'author': getattr(pdf_info, 'author', None),
                        'subject': getattr(pdf_info, 'subject', None),
                        'creator': getattr(pdf_info, 'creator', None),
                        'producer': getattr(pdf_info, 'producer', None)
                    })
                
                return metadata
                
        except Exception as e:
            logger.warning(f"Could not extract metadata from PDF {file_path}: {e}")
            return {'processor_type': 'pdf', 'extraction_error': str(e)}


class WordProcessor(BaseDocumentProcessor):
    """
    Process Microsoft Word documents (.docx) using python-docx.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.include_tables = self.config.get('include_tables', True)
    
    def _get_supported_extensions(self) -> List[str]:
        return ['.docx']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from Word document."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text if enabled
            if self.include_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(f"\\n[TABLE]\\n{table_text}\\n[/TABLE]\\n")
            
            if not text_parts:
                raise ProcessingError(f"No extractable text found in Word document: {file_path}")
            
            return "\\n".join(text_parts)
            
        except Exception as e:
            if isinstance(e, ProcessingError):
                raise
            raise ProcessingError(f"Failed to extract text from Word document {file_path}: {e}")
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a Word table."""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text if cell_text else "[empty]")
            table_data.append(" | ".join(row_data))
        return "\\n".join(table_data)
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        try:
            doc = DocxDocument(file_path)
            
            metadata = {
                'processor_type': 'word',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            # Extract document properties
            if doc.core_properties:
                props = doc.core_properties
                metadata.update({
                    'title': props.title,
                    'author': props.author,
                    'subject': props.subject,
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None
                })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Could not extract metadata from Word document {file_path}: {e}")
            return {'processor_type': 'word', 'extraction_error': str(e)}


class TextProcessor(BaseDocumentProcessor):
    """
    Process plain text files (.txt, .md, .py, etc.).
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.encodings = self.config.get('encodings', ['utf-8', 'ascii', 'latin-1'])
        self.max_file_size = self.config.get('max_file_size', 10 * 1024 * 1024)  # 10MB
    
    def _get_supported_extensions(self) -> List[str]:
        return ['.txt', '.md', '.py', '.js', '.json', '.csv', '.log']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from plain text file."""
        # Check file size
        file_size = Path(file_path).stat().st_size
        if file_size > self.max_file_size:
            raise ProcessingError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        # Try different encodings
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    
                    if not content.strip():
                        raise ProcessingError(f"File is empty: {file_path}")
                    
                    return content
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise ProcessingError(f"Failed to read text file {file_path}: {e}")
        
        raise ProcessingError(f"Could not decode text file {file_path} with any supported encoding")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from text file."""
        try:
            content = self.extract_text(file_path)
            lines = content.split('\\n')
            
            metadata = {
                'processor_type': 'text',
                'line_count': len(lines),
                'char_count_no_spaces': len(content.replace(' ', '')),
                'encoding_used': self._detect_encoding(file_path)
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Could not extract metadata from text file {file_path}: {e}")
            return {'processor_type': 'text', 'extraction_error': str(e)}
    
    def _detect_encoding(self, file_path: str) -> str:
        """Detect the encoding used to read the file."""
        for encoding in self.encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    file.read(100)  # Try to read first 100 chars
                return encoding
            except UnicodeDecodeError:
                continue
        return 'unknown'


class WebProcessor(BaseDocumentProcessor):
    """
    Process web pages and HTML files using BeautifulSoup.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.timeout = self.config.get('timeout', 30)
        self.user_agent = self.config.get('user_agent', 
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
    
    def _get_supported_extensions(self) -> List[str]:
        return ['.html', '.htm']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from HTML file or web URL."""
        try:
            # Check if it's a URL or file path
            if self._is_url(file_path):
                html_content = self._fetch_web_content(file_path)
            else:
                html_content = self._read_html_file(file_path)
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove unwanted elements
            self._clean_html(soup)
            
            # Extract text content
            text_content = soup.get_text(separator='\\n', strip=True)
            
            if not text_content.strip():
                raise ProcessingError(f"No extractable text found in HTML: {file_path}")
            
            return text_content
            
        except Exception as e:
            if isinstance(e, ProcessingError):
                raise
            raise ProcessingError(f"Failed to extract text from HTML {file_path}: {e}")
    
    def _is_url(self, path: str) -> bool:
        """Check if the path is a URL."""
        return path.startswith(('http://', 'https://'))
    
    def _fetch_web_content(self, url: str) -> str:
        """Fetch content from web URL."""
        headers = {'User-Agent': self.user_agent}
        
        try:
            response = requests.get(url, headers=headers, timeout=self.timeout)
            response.raise_for_status()
            return response.text
            
        except requests.RequestException as e:
            raise ProcessingError(f"Failed to fetch web content from {url}: {e}")
    
    def _read_html_file(self, file_path: str) -> str:
        """Read HTML content from local file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _clean_html(self, soup: BeautifulSoup) -> None:
        """Remove unwanted HTML elements."""
        # Remove scripts, styles, and other non-content elements
        unwanted_tags = ['script', 'style', 'nav', 'header', 'footer']
        
        for tag_name in unwanted_tags:
            for element in soup.find_all(tag_name):
                element.decompose()
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from HTML file or web page."""
        try:
            if self._is_url(file_path):
                html_content = self._fetch_web_content(file_path)
                source_type = 'web_url'
            else:
                html_content = self._read_html_file(file_path)
                source_type = 'html_file'
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            metadata = {
                'processor_type': 'web',
                'source_type': source_type,
                'url': file_path if self._is_url(file_path) else None
            }
            
            # Extract basic HTML metadata
            if soup.title:
                metadata['title'] = soup.title.string.strip() if soup.title.string else None
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Could not extract metadata from HTML {file_path}: {e}")
            return {'processor_type': 'web', 'extraction_error': str(e)}


class DocumentProcessorFactory:
    """
    Factory class for creating appropriate document processors based on file type.
    """
    
    def __init__(self):
        self.processors = {
            '.pdf': PDFProcessor,
            '.docx': WordProcessor,
            '.txt': TextProcessor,
            '.md': TextProcessor,
            '.py': TextProcessor,
            '.js': TextProcessor,
            '.json': TextProcessor,
            '.csv': TextProcessor,
            '.log': TextProcessor,
            '.html': WebProcessor,
            '.htm': WebProcessor
        }
    
    def create_processor(self, file_path: str, config: Optional[Dict[str, Any]] = None) -> BaseDocumentProcessor:
        """
        Create appropriate processor for the given file.
        
        Args:
            file_path: Path to the file or URL
            config: Optional configuration for the processor
            
        Returns:
            Appropriate document processor instance
            
        Raises:
            UnsupportedFormatError: If file format is not supported
        """
        # Handle URLs
        if file_path.startswith(('http://', 'https://')):
            return WebProcessor(config)
        
        # Handle file extensions
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.processors:
            raise UnsupportedFormatError(f"Unsupported file format: {file_extension}")
        
        processor_class = self.processors[file_extension]
        return processor_class(config)
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of all supported file extensions."""
        return list(self.processors.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        if file_path.startswith(('http://', 'https://')):
            return True
        
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.processors


# Example usage and testing
if __name__ == "__main__":
    factory = DocumentProcessorFactory()
    
    print("Supported file extensions:")
    print(factory.get_supported_extensions())
    
    # Test with different file types
    test_files = ["test.pdf", "test.docx", "test.txt", "https://example.com"]
    
    for test_file in test_files:
        try:
            processor = factory.create_processor(test_file)
            print(f"✅ {test_file} -> {type(processor).__name__}")
        except UnsupportedFormatError as e:
            print(f"❌ {test_file} -> {e}")